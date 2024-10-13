from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from src.lab2.museum.virtual_museum import VirtualMuseum
from src.lab2.reports.report_maker import ReportMaker
import src.lab2.client as client


class WordColumnStyle:
    def __init__(self, index: int, header: str, width: Inches) -> None:
        self.__index = index
        self.__header = header
        self.__width = width

    @property
    def index(self) -> int:
        return self.__index

    @property
    def header(self) -> str:
        return self.__header

    @property
    def width(self) -> Inches:
        return self.__width


class WordReportMaker(ReportMaker):
    __TABLE_ALIGNMENT = WD_TABLE_ALIGNMENT.CENTER
    __ALIGNMENT = WD_ALIGN_PARAGRAPH.RIGHT
    __INIT_ROW = 0
    __COLUMNS = (
        WordColumnStyle(0, 'Id', Inches(0.75)),
        WordColumnStyle(1, 'Название', Inches(2)),
        WordColumnStyle(2, 'Описание', Inches(3)),
        WordColumnStyle(3, 'Тип', Inches(1)),
        WordColumnStyle(4, 'Id коллекции', Inches(1)),
        WordColumnStyle(5, 'Просмотров пользователя', Inches(1)),
        WordColumnStyle(6, 'Всего просмотров', Inches(1)),
    )

    def __init__(self, museum: VirtualMuseum) -> None:
        super().__init__(museum)
        self.__doc = Document()
        self.__table = None
        self.__row = WordReportMaker.__INIT_ROW

    def make(self, path: str = None) -> None:
        self.__preprocess()
        self.__write_data()
        self.__postprocess()
        path = path or f"user_{self._museum.user.id}_{self._museum.user.name}.docx"
        self.__doc.save(path)

    def __preprocess(self) -> None:
        self.__table = self.__doc.add_table(rows=1, cols=len(WordReportMaker.__COLUMNS))
        self.__table.alignment = WordReportMaker.__TABLE_ALIGNMENT
        self.__table.autofit = False

        for i, column in enumerate(WordReportMaker.__COLUMNS):
            self.__table.cell(0, i).text = column.header

    def __postprocess(self) -> None:
        for i in range(self.__row + 1):
            for j, column in enumerate(WordReportMaker.__COLUMNS):
                self.__table.cell(i, j).width = column.width
                self.__table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for section in self.__doc.sections:
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.orientation = WD_ORIENT.LANDSCAPE

    def __write_data(self):
        data = client.get_user_info(self._museum.user.id)
        for row in data:
            self.__row += 1
            self.__table.add_row()
            self.__table.cell(self.__row, 0).text = str(row.get('item_id', None))
            self.__table.cell(self.__row, 1).text = str(row.get('title', None))
            self.__table.cell(self.__row, 2).text = str(row.get('descr', None))
            self.__table.cell(self.__row, 3).text = str(row.get('type', None))
            self.__table.cell(self.__row, 4).text = str(row.get('collection_id', None))
            self.__table.cell(self.__row, 5).text = str(row.get('user_views', None))
            self.__table.cell(self.__row, 6).text = str(row.get('total_views', None))
